import io
import json
import traceback
from unittest import result
from django.shortcuts import get_object_or_404, render, redirect
from django.urls import reverse
from httplib2 import Authentication
from jsonschema import ValidationError
from .models import *
from django.db import connection
from django.contrib import messages
from django.db.models import Sum, Value
from django.db.models.functions import Coalesce
import logging
from django.http import HttpResponse, HttpResponseServerError, JsonResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import authenticate, login as auth_login
from django.shortcuts import render, redirect
from django.contrib.auth import login
from django.contrib.auth.models import User
from docx import Document
from django.core.exceptions import ObjectDoesNotExist
from docx.shared import Inches





def index(request):
    # Получаем информацию о сотруднике из сессии, если она есть
    employee_id = request.session.get('employee_id')

    # Добавляем остальные данные, которые вы используете в шаблоне
    warehouses = Warehouse.objects.all()
    offices = Office.objects.all()
    employees = Employee.objects.all()
    products = Product.objects.all()
    orders = Order.objects.all()
    components = Components.objects.all()
    warehouse_movements = WarehouseMovement.objects.all()
    products_movements = ProductsMovement.objects.all()
    categorys = Category.objects.all() 

    context = {
        'employee_id': employee_id,
        'warehouses': warehouses,
        'offices': offices,
        'employees': employees,
        'products': products,
        'orders': orders,
        'components': components,
        'warehouse_movements': warehouse_movements,
        'products_movements': products_movements,
        'categorys' : categorys,
    }

    return render(request, 'index.html', context)

#-------------------------------------------------------------------------------------

def category(request):
    categorys = Category.objects.all()
    return render(request, 'category.html', {
        'categorys' : categorys
    })

def add_category(request):
    if request.method == 'POST':
        name = request.POST['name']

        new_category = Category(name=name)
        new_category.save()
        return redirect('category')

    return render(request, 'add_category.html')

#-------------------------------------------------------------------------------------

def warehouse(request):
    warehouses = Warehouse.objects.all()
    return render(request, 'warehouse.html', {
        'warehouses': warehouses
    })

def update_hidden_status(request):
    if request.method == 'POST':
        warehouse_id = request.POST.get('warehouse_id')
        is_hidden_str = request.POST.get('is_hidden')

        # Convert the string "true" to a boolean value
        is_hidden = is_hidden_str.lower() == 'true'

        try:
            warehouse = Warehouse.objects.get(pk=warehouse_id)
            warehouse.hidden = is_hidden
            warehouse.save()
            return JsonResponse({'success': True, 'hidden': warehouse.hidden})
        except Warehouse.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Warehouse not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

def add_warehouse(request):
    if request.method == 'POST':
        address = request.POST['address']
        phone = request.POST['phone']

        new_warehouse = Warehouse(address=address, phone=phone)
        new_warehouse.save()
        return redirect('warehouse')
    return render(request, 'add_warehouse.html')

def edit_warehouse(request, warehouse_id):
    try:
        warehouse = get_object_or_404(Warehouse, pk=warehouse_id)

        if request.method == 'POST':
            # Обработка отправки формы для обновления деталей склада
            address = request.POST['address']
            phone = request.POST['phone']

            warehouse.address = address
            warehouse.phone = phone
            warehouse.save()

            return redirect('warehouse')

        return render(request, 'edit/edit_warehouse.html', {'warehouse': warehouse})
    except Exception as e:
        # Отображение информации об ошибке
        return HttpResponseServerError(f"Internal Server Error: {str(e)}")
    
def export_to_word(request):
    warehouses = Warehouse.objects.all()

    warehouses = warehouses.filter(hidden=False)

    document = Document()
    document.add_heading('Warehouse Export', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add table header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Address'
    header_cells[1].text = 'Phone'

    # Add data to the table
    for warehouse in warehouses:
        row_cells = table.add_row().cells
        row_cells[0].text = warehouse.address
        row_cells[1].text = warehouse.phone

    # Response with Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=warehouse_export.docx'
    document.save(response)

    return response
#-------------------------------------------------------------------------------------

def product(request):
    products = Product.objects.all()
    return render(request, 'products.html', {
        'products': products,
    })

def add_products(request):
    categories = Category.objects.all()  # Fetch all categories

    if request.method == 'POST':
        name = request.POST['name']
        category_id = request.POST['category']

        new_product = Product(name=name, category_id=category_id)
        new_product.save()
        return redirect('products')

    return render(request, 'add_products.html', {'categories': categories})

def update_hidden_status_products(request):
    if request.method == 'POST':
        products_id = request.POST.get('product_id')  # Corrected to 'products_id'
        is_hidden_str = request.POST.get('is_hidden')

        # Convert the string "true" to a boolean value
        is_hidden = is_hidden_str.lower() == 'true'

        try:
            product = Product.objects.get(pk=products_id)  # Corrected to Product
            product.hidden = is_hidden
            product.save()
            return JsonResponse({'success': True, 'hidden': product.hidden})
        except Product.DoesNotExist:  # Corrected to Product.DoesNotExist
            return JsonResponse({'success': False, 'error': 'Product not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})

def edit_products(request, product_id):
    try:
        product = get_object_or_404(Product, pk=product_id)
        categories = Category.objects.all()  # Получение всех категорий

        if request.method == 'POST':
            # Обработка отправки формы для обновления деталей продукта
            name = request.POST['name']
            category_id = request.POST['category']

            product.name = name
            product.category_id = category_id
            product.save()

            return redirect('products')

        return render(request, 'edit/edit_products.html', {'product': product, 'categories': categories})
    except Exception as e:
        return HttpResponseServerError(f"Internal Server Error: {str(e)}")

def export_to_word_products(request):
    products = Product.objects.all()

    products = products.filter(hidden=False)

    document = Document()
    document.add_heading('Products Export', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Category'

    for product in products:
        row_cells = table.add_row().cells
        row_cells[0].text = product.name
        row_cells[1].text = product.category.name  # Assuming category has a 'name' field

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=products_export.docx'
    document.save(response)

    return response
#-------------------------------------------------------------------------------------

def offices(request):
    offices = Office.objects.all()
    return render(request, 'offices.html', {
        'offices': offices
    })

def add_offices(request):
    if request.method == 'POST':
        address = request.POST['address']
        area = request.POST['area']
        phone = request.POST['phone']

        new_office = Office.objects.create(address=address, area=area, phone=phone)
        return redirect('offices')
    return render(request, 'add_offices.html')

def edit_office(request, office_id):
    office = get_object_or_404(Office, pk=office_id)

    if request.method == 'POST':
        address = request.POST['address']
        area = request.POST['area']
        phone = request.POST['phone']

        office.address = address
        office.area = area
        office.phone = phone
        office.save()

        return redirect('offices')

    return render(request, 'edit/edit_office.html', {'office': office})

def update_hidden_status_offices(request):
    if request.method == 'POST':
        office_id = request.POST.get('office_id')
        is_hidden_str = request.POST.get('is_hidden')

        is_hidden = is_hidden_str.lower() == 'true'

        try:
            office = Office.objects.get(pk=office_id)
            office.hidden = is_hidden
            office.save()
            return JsonResponse({'success': True, 'hidden': office.hidden})
        except Office.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Office not found'})
        except ValueError as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})

def export_to_word_office(request):
    offices = Office.objects.all()

    offices = offices.filter(hidden=False)

    document = Document()
    document.add_heading('Offices Export', level=1)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Address'
    header_cells[1].text = 'Area'
    header_cells[2].text = 'Phone'

    for office in offices:
        row_cells = table.add_row().cells
        row_cells[0].text = office.address
        row_cells[1].text = str(office.area)  # Ensure area is converted to string
        row_cells[2].text = office.phone

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=office_export.docx'
    document.save(response)

    return response
#-------------------------------------------------------------------------------------

def employees(request):
    employees_list = Employee.objects.all()
    return render(request, 'employees.html', {
        'employees': employees_list
    })

def add_employees(request):
    if request.method == 'POST':
        name = request.POST['name']
        last_name = request.POST['last_name']

        new_employee = Employee(name=name, lastName=last_name)
        new_employee.save()
        return redirect('employees')
    
    return render(request, 'add_employees.html')

def update_hidden_status_employees(request):
    if request.method == 'POST':
        employees_id = request.POST.get('employee_id')
        is_hidden_str = request.POST.get('is_hidden')

        # Convert the string "true" to a boolean value
        is_hidden = is_hidden_str.lower() == 'true'

        try:
            employee = Employee.objects.get(pk=employees_id)
            employee.hidden = is_hidden
            employee.save()
            return JsonResponse({'success': True, 'hidden': employee.hidden})
        except Employee.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Employee not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

def edit_employees(request, employee_id):
    try:
        employee = get_object_or_404(Employee, pk=employee_id)

        if request.method == 'POST':
            name = request.POST['name']
            last_name = request.POST['last_name']

            employee.name = name
            employee.lastName = last_name
            employee.save()

            return redirect('employees')

        return render(request, 'edit/edit_employees.html', {'employee': employee})
    except Exception as e:
        # Displaying error information
        return HttpResponseServerError(f"Internal Server Error: {str(e)}")
    
def login(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        last_name = request.POST.get('last_name')

        try:
            employee = Employee.objects.get(name=name, lastName=last_name)
            request.session['employee_id'] = employee.id
            request.session['is_admin'] = employee.is_admin
            return redirect('index')
        except Employee.DoesNotExist:
            pass  # Employee not found

        messages.error(request, 'Incorrect name or last name of the employee.')

    return render(request, 'login.html')

def export_to_word_employees(request):
    employees = Employee.objects.all()

    employees = employees.filter(hidden=False)

    document = Document()
    document.add_heading('Employees Export', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Last Name'

    for employee in employees:
        row_cells = table.add_row().cells
        row_cells[0].text = employee.name
        row_cells[1].text = employee.lastName

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=employees_export.docx'
    document.save(response)

    return response

#-------------------------------------------------------------------------------------

def orders(request):
    orders = Order.objects.all()
    return render(request, 'orders.html', {'orders': orders})

def add_orders(request):
    employees = Employee.objects.all()
    offices = Office.objects.all()
    products = Product.objects.all()

    if request.method == 'POST':
        date = request.POST.get('date')
        status = request.POST.get('status')
        employee_id = request.POST.get('employee')
        product_id = request.POST.get('product')
        comment = request.POST.get('comment')
        office_id = request.POST.get('office')

        employee = Employee.objects.get(pk=employee_id)
        product = Product.objects.get(pk=product_id)
        office = Office.objects.get(pk=office_id)

        order = Order(
            date=date,
            status=status,
            employee=employee,
            product=product,
            comment=comment,
            IdOffice=office,
        )

        order.save()

        return redirect('orders')

    return render(request, 'add_orders.html', {'employees': employees, 'offices': offices, 'products': products})

def edit_orders(request, order_id):
    try:
        order = get_object_or_404(Order, pk=order_id)
        employees = Employee.objects.all()
        offices = Office.objects.all()
        products = Product.objects.all()

        if request.method == 'POST':
            date = request.POST.get('date')
            status = request.POST.get('status')
            employee_id = request.POST.get('employee')
            product_id = request.POST.get('product')
            comment = request.POST.get('comment')
            office_id = request.POST.get('office')

            employee = Employee.objects.get(pk=employee_id)
            product = Product.objects.get(pk=product_id)
            office = Office.objects.get(pk=office_id)

            order.date = date
            order.status = status
            order.employee = employee
            order.product = product
            order.comment = comment
            order.IdOffice = office
            order.save()

            return redirect('orders')

        return render(request, 'edit/edit_orders.html', {'order': order, 'employees': employees, 'offices': offices, 'products': products})
    except Exception as e:
        # Отображение информации об ошибке
        return HttpResponseServerError(f"Internal Server Error: {str(e)}")

def update_hidden_status_orders(request):
    if request.method == 'POST':
        order_id = request.POST.get('order_id')
        is_checked = request.POST.get('is_checked')

        try:
            # Получение объекта заказа по его идентификатору
            order = Order.objects.get(pk=order_id)
            # Преобразование строки "true" в логическое значение
            order.hidden = is_checked.lower() == 'true'
            order.save()
            return JsonResponse({'success': True, 'hidden': order.hidden})
        except ObjectDoesNotExist:
            return JsonResponse({'success': False, 'error': 'Order not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

def export_to_word_orders(request):
    selected_status = request.GET.get('status')

    # Filtering orders based on selected status
    if selected_status:
        orders = Order.objects.filter(status=selected_status)
    else:
        orders = Order.objects.all()
    
    # Excluding hidden orders
    orders = orders.filter(hidden=False)

    document = Document()
    document.add_heading('Exported Orders', level=1)

    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Date'
    header_cells[1].text = 'Status'
    header_cells[2].text = 'Employee'
    header_cells[3].text = 'Product'
    header_cells[4].text = 'Comment'
    header_cells[5].text = 'Office'

    for order in orders:
        row_cells = table.add_row().cells
        row_cells[0].text = str(order.date)
        row_cells[1].text = order.status
        row_cells[2].text = order.employee.name
        row_cells[3].text = order.product.name
        row_cells[4].text = order.comment
        row_cells[5].text = order.IdOffice.address

    for col in table.columns:
        col.width = Inches(2)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=orders_export.docx'
    document.save(response)

    return response


#-------------------------------------------------------------------------------------

def components(request):
    components = Components.objects.all()
    return render(request, 'components.html', {
        'components': components
    })

def add_components(request):
    if request.method == 'POST':
        name = request.POST['name']

        new_component = Components(name=name)
        new_component.save()
        return redirect('components')
    return render(request, 'add_components.html')
#-------------------------------------------------------------------------------------

def warehouse_movements(request):
    warehouse_movements = WarehouseMovement.objects.all()
    return render(request, 'warehouse_movements.html', {
        'warehouse_movements': warehouse_movements
    })

logger = logging.getLogger(__name__)

def calculate_total_quantity(request):
    try:
        selected_component_id = request.GET.get('component_id')
        selected_warehouse_minus_id = request.GET.get('warehouse_minus_id')

        if not selected_component_id or not selected_warehouse_minus_id:
            return JsonResponse({'error': 'Invalid input. Please provide both component_id and warehouse_minus_id.'}, status=400)

        warehouse_movement_quantity_minus = WarehouseMovement.objects.filter(
            IdComponents_id=selected_component_id,
            IdWarehouseMinus_id=selected_warehouse_minus_id
        ).aggregate(
            warehouse_movement_quantity_minus=Sum('quantity')
        )['warehouse_movement_quantity_minus'] or 0

        warehouse_movement_quantity_plus = WarehouseMovement.objects.filter(
            IdComponents_id=selected_component_id,
            IdWarehousePlus_id=selected_warehouse_minus_id
        ).aggregate(
            warehouse_movement_quantity_plus=Sum('quantity')
        )['warehouse_movement_quantity_plus'] or 0

        products_movement_quantity = ProductsMovement.objects.filter(
            IdComponents_id=selected_component_id,
            IdWarehouse__id=selected_warehouse_minus_id,
            status='updated'
        ).aggregate(
            products_movement_quantity=Sum('quantity')
        )['products_movement_quantity'] or 0

        substitution_quantity = ProductsMovement.objects.filter(
            IdComponents_id=selected_component_id,
            IdWarehouse__id=selected_warehouse_minus_id,
            status='substitution'
        ).aggregate(
            substitution_quantity=Sum('quantity')
        )['substitution_quantity'] or 0

        total_quantity = warehouse_movement_quantity_plus - products_movement_quantity - warehouse_movement_quantity_minus + substitution_quantity

        warehouses_with_component = Warehouse.objects.filter(
            warehouse_plus__IdComponents_id=selected_component_id
        ).annotate(
            total_quantity=Coalesce(Sum('warehouse_plus__quantity'), Value(0)) - Coalesce(Sum('warehouse_minus__quantity'), Value(0))
        )

        # Преобразование QuerySet в список словарей
        warehouses_with_component_list = list(warehouses_with_component.values())

        return JsonResponse({'total_quantity': total_quantity or 0, 'warehouses_with_component': warehouses_with_component_list})
    except Exception as e:
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)

    
def add_warehouse_movement(request):
    total_quantity = 0
    warehouses_with_component = []

    if request.method == 'POST':
        try:
            selected_component_id = request.POST.get('component_name')
            quantity_raw = request.POST.get('quantity')
            quantity = int(quantity_raw) if quantity_raw.strip() else 0
            id_warehouse_plus = request.POST.get('id_warehouse_plus')
            id_warehouse_minus = request.POST.get('id_warehouse_minus')
            date_str = request.POST.get('date')
            comment = request.POST.get('comment')



            selected_component = get_object_or_404(Components, id=selected_component_id)

            # Check if warehouse_plus_id is the same as warehouse_minus_id
            if id_warehouse_plus == id_warehouse_minus:
                return HttpResponse('Warehouse plus and minus cannot be the same.')

            # Convert the date string to a timezone-aware object
            current_timezone = timezone.get_current_timezone()
            date = current_timezone.localize(timezone.datetime.strptime(date_str, '%Y-%m-%dT%H:%M'))

            # Check if IdWarehouseMinus_id is present before saving
            if id_warehouse_minus is None:
                return HttpResponse('IdWarehouseMinus_id cannot be empty.')

            # Add print statement for debugging purposes
            print('IdWarehouseMinus_id:', id_warehouse_minus)

            # Check if IdWarehouseMinus_id is present before saving (additional check)
            if not id_warehouse_minus:
                return HttpResponse('IdWarehouseMinus_id cannot be empty.')

            # Check if the quantity is non-negative and less than or equal to available quantity
            if quantity <= 0:
                return HttpResponse('Quantity must be a non-negative value.')



            # Create and save the warehouse movement object
            warehouse_movement = WarehouseMovement(
                IdComponents=selected_component,
                quantity=quantity,
                IdWarehousePlus_id=id_warehouse_plus,
                IdWarehouseMinus_id=id_warehouse_minus,
                date=date,
                Comment=comment,
            )
            warehouse_movement.save()

        except (Components.DoesNotExist, ValueError) as e:
            # Return a more informative error message
            return HttpResponse(f'Input error. Please check your data. Error: {e}')

    context = {
        'warehouses': Warehouse.objects.all(),
        'components': Components.objects.all(),
        'total_quantity': total_quantity,
        'warehouses_with_component': warehouses_with_component,
        'warehouse_movements': WarehouseMovement.objects.all(),
    }

    return render(request, 'add_warehouse_movements.html', context)

def export_warehouse_movements(request):
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')

    warehouse_movements = WarehouseMovement.objects.all()

    if start_date and end_date:
        warehouse_movements = warehouse_movements.filter(date__range=[start_date, end_date])

    # Создание документа Word
    document = Document()
    document.add_heading('Warehouse Movements', 0)
    
    # Добавление данных в документ
    for warehouse_movement in warehouse_movements:
        document.add_paragraph(f"Comment: {warehouse_movement.Comment}")
        document.add_paragraph(f"Quantity: {warehouse_movement.quantity}")
        document.add_paragraph(f"Components: {warehouse_movement.IdComponents}")
        document.add_paragraph(f"IdWarehousePlus: {warehouse_movement.IdWarehousePlus}")
        document.add_paragraph(f"IdWarehouseMinus: {warehouse_movement.IdWarehouseMinus}")
        document.add_paragraph(f"Date: {warehouse_movement.date}")
        document.add_paragraph("-------------------")

    # Создание буфера для хранения данных документа
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Отправка файла пользователю
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=WarehouseMovements.docx'
    return response
    
#-------------------------------------------------------------------------------------

def products_movements(request):
    products_movements = ProductsMovement.objects.all()
    return render(request, 'products_movements.html', {
        'products_movements': products_movements
    })

def add_products_movements(request):
    if request.method == 'POST':
        id_warehouse = request.POST.get('IdWarehouse')
        quantity = request.POST.get('quantity')
        id_components = request.POST.get('IdComponents')
        id_product = request.POST.get('IdProduct')
        status = request.POST.get('status')

        products_movement = ProductsMovement.objects.create(
            IdWarehouse=Warehouse.objects.get(id=id_warehouse),
            quantity=quantity,
            IdComponents=Components.objects.get(id=id_components),
            IdProduct=Product.objects.get(id=id_product),
            status=status
        )

        return redirect('products_movements')

    warehouses = Warehouse.objects.all()
    components = Components.objects.all()
    products = Product.objects.all()

    current_datetime = ProductsMovement._meta.get_field('datetime').default()

    return render(request, 'add_products_movements.html', {'warehouses': warehouses, 'components': components, 'products': products})
#-------------------------------------------------------------------------------------